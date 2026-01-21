"""Export helpers.

Keep all file-export logic (DOCX, etc.) in one place so pages can share it.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt


def create_docx_from_markdown(text: str, *, title: Optional[str] = None) -> BytesIO:
    """Convert a markdown-ish text blob into a DOCX.

    Supports:
      - Markdown headings starting with ##, ###, #### (converted to Word headings)
      - Strips **bold** markers (keeps the text)

    This is intentionally simple and robust for marketing copy exports.
    """

    doc = Document()

    # Set sensible defaults
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    except Exception:
        pass

    if title and title.strip():
        doc.add_heading(title.strip(), level=1)

    for raw_line in (text or "").splitlines():
        line = (raw_line or "").rstrip()
        if not line.strip():
            # Preserve spacing with a blank paragraph
            doc.add_paragraph("")
            continue

        m = re.match(r"^(#{2,6})\s+(.*)$", line)
        if m:
            hashes = m.group(1)
            heading = m.group(2).strip()
            # Map markdown ##..###### to Word heading levels 2..6
            level = min(max(len(hashes), 2), 6)
            # Word levels start at 1. We treat ## as level 2.
            doc.add_heading(heading.replace("**", ""), level=level)
            continue

        clean = line.replace("**", "")
        doc.add_paragraph(clean)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
